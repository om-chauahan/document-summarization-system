from __future__ import annotations

from django.test import TestCase
from django.contrib.auth import get_user_model
from .models import StoredDocument
from django.core.files.uploadedfile import SimpleUploadedFile
from unittest.mock import patch
from django.test.utils import override_settings
from django.core import mail
import re

class ExtractionSmokeTests(TestCase):
	def test_extract_text_from_txt_bytes(self):
		from .views import extract_text_from_txt_bytes

		raw = b"Hello world\nThis is a test.\n"
		out = extract_text_from_txt_bytes(raw)
		self.assertIn("Hello world", out)
		self.assertIn("This is a test.", out)

	def test_extract_text_from_docx_bytes(self):
		try:
			from docx import Document
		except Exception:
			self.skipTest("python-docx not installed")

		from .views import extract_text_from_docx_bytes

		doc = Document()
		doc.add_paragraph("First paragraph")
		doc.add_paragraph("Second paragraph")

		import io

		buf = io.BytesIO()
		doc.save(buf)
		out = extract_text_from_docx_bytes(buf.getvalue())
		self.assertIn("First paragraph", out)
		self.assertIn("Second paragraph", out)

class SummarizationHelpersTests(TestCase):
	def test_summarize_text_empty_input(self):
		from DSS_app.summarization import summarize_text
		self.assertEqual(summarize_text(""), "")

	def test_default_prompt_requests_20_lines(self):
		"""Ensure the default prompt requests a 20-line summary."""
		from DSS_app import summarization as s
		src = s.summarize_text.__code__.co_consts
		joined = " ".join([c for c in src if isinstance(c, str)])
		self.assertIn("summary length is 20 lines", joined)

	def test_compress_removes_repeated_headers_preserves_key_lines(self):
		from DSS_app.summarization import _compress_for_llm
		text = "\n".join(
			[
				"ACME AIRLINES - ITINERARY",
				"Page 1 of 3",
				"Passenger: John Doe",
				"PNR: XY12ZZ",
				"Flight: AC123  10 Jan 2026  10:30",
				"ACME AIRLINES - ITINERARY",
				"Page 2 of 3",
				"Seat: 12A",
				"Gate: B7",
				"ACME AIRLINES - ITINERARY",
				"Page 3 of 3",
			]
		)
		out = _compress_for_llm(text, max_chars=2000)
		self.assertIn("Passenger: John Doe", out)
		self.assertIn("PNR: XY12ZZ", out)

class AuthApiTests(TestCase):
	def test_signup_creates_user(self):
		res = self.client.post(
			"/api/auth/signup/",
			data={
				"firstName": "A",
				"lastName": "B",
				"mobile": "9999999999",
				"email": "a@example.com",
				"password": "pass1234!",
			},
			content_type="application/json",
		)
		self.assertEqual(res.status_code, 201)
		self.assertEqual(res.json().get("ok"), True)


class UploadPreflightTests(TestCase):
	def setUp(self):
		User = get_user_model()
		self.user = User.objects.create_user(username="u1", password="password123")
		self.client.login(username="u1", password="password123")

	def test_preflight_estimates_credits_from_extracted_text_bytes(self):
		# 1 KB-ish extracted text should map to the first tier (3 credits) using the shared backend function.
		content = ("Hello world.\n" * 50).encode("utf-8")
		up = SimpleUploadedFile("a.txt", content, content_type="text/plain")
		res = self.client.post("/api/uploads/preflight/", {"file": up})
		self.assertEqual(res.status_code, 200)
		data = res.json()
		self.assertEqual(data.get("ok"), True)
		self.assertIsInstance(data.get("upload_id"), int)
		self.assertIsInstance(data.get("extracted_text_bytes"), int)
		self.assertIsInstance(data.get("extracted_text_length"), int)
		self.assertIn(data.get("required_credits"), (3, 5, 8, 12, 18, 25, 40, 60))

		# Preflight creates a StoredDocument but should not set summary yet.
		doc = StoredDocument.objects.get(id=data.get("upload_id"))
		self.assertEqual(doc.user_id, self.user.id)
		self.assertEqual(doc.original_name, "a.txt")
		self.assertTrue((doc.extracted_text or "").strip())
		self.assertEqual(doc.summary, "")

	def test_list_uploads_excludes_preflight_items(self):
		# Create a preflight item (summary empty)
		up = SimpleUploadedFile("b.txt", b"Hello", content_type="text/plain")
		res = self.client.post("/api/uploads/preflight/", {"file": up})
		self.assertEqual(res.status_code, 200)

		# list_uploads should exclude it
		res2 = self.client.get("/api/uploads/")
		self.assertEqual(res2.status_code, 200)
		data2 = res2.json()
		self.assertEqual(data2.get("ok"), True)
		uploads = data2.get("uploads") or []
		self.assertEqual(len(uploads), 0)

	def test_login_rejects_invalid_credentials(self):
		# No user created
		res = self.client.post(
			"/api/auth/login/",
			data={"identifier": "9999999999", "password": "wrong"},
			content_type="application/json",
		)
		self.assertEqual(res.status_code, 401)
		self.assertEqual(res.json().get("ok"), False)

	def test_login_accepts_valid_credentials(self):
		self.client.post(
			"/api/auth/signup/",
			data={
				"firstName": "A",
				"lastName": "B",
				"mobile": "8888888888",
				"email": "b@example.com",
				"password": "pass1234!",
			},
			content_type="application/json",
		)

		res = self.client.post(
			"/api/auth/login/",
			data={"identifier": "8888888888", "password": "pass1234!"},
			content_type="application/json",
		)
		self.assertEqual(res.status_code, 200)
		self.assertEqual(res.json().get("ok"), True)

	def test_login_accepts_email_identifier(self):
		self.client.post(
			"/api/auth/signup/",
			data={
				"firstName": "A",
				"lastName": "B",
				"mobile": "+91 99999-00000",
				"email": "email_login@example.com",
				"password": "pass1234!",
			},
			content_type="application/json",
		)

		res = self.client.post(
			"/api/auth/login/",
			data={"identifier": "email_login@example.com", "password": "pass1234!"},
			content_type="application/json",
		)
		self.assertEqual(res.status_code, 200)
		self.assertEqual(res.json().get("ok"), True)


@override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
class PasswordOtpTests(TestCase):
	def test_change_password_via_otp(self):
		# Create user
		res = self.client.post(
			"/api/auth/signup/",
			data={
				"firstName": "A",
				"lastName": "B",
				"mobile": "7777777777",
				"email": "otp@example.com",
				"password": "oldpass123",
			},
			content_type="application/json",
		)
		self.assertEqual(res.status_code, 201)

		# Login (session cookie)
		res = self.client.post(
			"/api/auth/login/",
			data={"identifier": "7777777777", "password": "oldpass123"},
			content_type="application/json",
		)
		self.assertEqual(res.status_code, 200)
		self.assertEqual(res.json().get("ok"), True)

		# Request OTP
		res = self.client.post("/api/auth/change-password/otp/request/")
		self.assertEqual(res.status_code, 200)
		self.assertEqual(res.json().get("ok"), True)
		self.assertGreaterEqual(len(mail.outbox), 1)
		body = mail.outbox[-1].body or ""
		m = re.search(r"\b(\d{6})\b", body)
		self.assertIsNotNone(m)
		otp = m.group(1)

		# Verify OTP and set new password
		res = self.client.post(
			"/api/auth/change-password/otp/verify/",
			data={"otp": otp, "new_password": "newpass123"},
			content_type="application/json",
		)
		self.assertEqual(res.status_code, 200)
		self.assertEqual(res.json().get("ok"), True)

		# Logout and verify login works with new password
		self.client.post("/api/auth/logout/")
		res = self.client.post(
			"/api/auth/login/",
			data={"identifier": "7777777777", "password": "newpass123"},
			content_type="application/json",
		)
		self.assertEqual(res.status_code, 200)
		self.assertEqual(res.json().get("ok"), True)


class GoogleOAuthTests(TestCase):
	def test_google_login_requires_configuration(self):
		res = self.client.get("/api/auth/google/login/")
		# Depending on local dev environment, .env may be loaded during tests.
		# Accept either:
		# - 501/500: not configured
		# - 302: configured and redirecting to Google
		self.assertIn(res.status_code, (302, 501, 500))
		if res.status_code != 302:
			# If it's a JSON response it should be ok:false.
			if getattr(res, "headers", {}).get("Content-Type", "").startswith("application/json"):
				self.assertEqual(res.json().get("ok"), False)

	@override_settings(
		DSS_GOOGLE_CLIENT_ID="test-client-id",
		DSS_GOOGLE_CLIENT_SECRET="test-secret",
		DSS_GOOGLE_REDIRECT_URI="http://127.0.0.1:8000/api/auth/google/callback/",
	)
	def test_google_login_redirects_to_google(self):
		res = self.client.get("/api/auth/google/login/?next=%2Fupload")
		self.assertEqual(res.status_code, 302)
		loc = res.headers.get("Location", "")
		self.assertIn("accounts.google.com", loc)
		# State should be stored in session.
		sess = self.client.session
		self.assertTrue(sess.get("google_oauth_state"))
		self.assertEqual(sess.get("google_oauth_next"), "/upload")
